"""End-to-end browser test: every tab, every viewport, every export.

Downloads are captured to ./_testdl and then opened with openpyxl / python-docx /
json / xml so a file that merely *downloads* but does not *open* still fails.
"""
import asyncio, json, os, pathlib, shutil, sys, zipfile
import xml.etree.ElementTree as ET
from playwright.async_api import async_playwright

DL = pathlib.Path(__file__).resolve().parent.parent / '_testdl'
FAILS, NOTES = [], []


def bad(m):
    FAILS.append(m)
    print('  FAIL ', m)


def good(m):
    print('  ok   ', m)


def note(m):
    NOTES.append(m)
    print('  note ', m)


async def main():
    if DL.exists():
        shutil.rmtree(DL)
    DL.mkdir()
    root = pathlib.Path(__file__).resolve().parent.parent
    uri = (root / 'docs' / 'index.html').as_uri()
    # read the tab list from the page itself, so a new tab is covered automatically
    # and a removed one cannot leave a stale test passing
    import re as _re
    tabs = _re.findall(r'id="tab-([a-z]+)"',
                       (pathlib.Path(__file__).resolve().parent.parent
                        / 'docs' / 'index.html').read_text(encoding='utf-8'))

    async with async_playwright() as p:
        b = await p.chromium.launch()

        # ---------------- console / layout across viewports ----------------
        print('\n[A] tabs, console errors, horizontal overflow')
        for w, h, scheme in [(1440, 1000, 'light'), (1320, 1000, 'dark'),
                             (820, 900, 'light'), (400, 900, 'light')]:
            pg = await b.new_page(viewport={'width': w, 'height': h}, color_scheme=scheme)
            errs = []
            pg.on('pageerror', lambda e: errs.append(str(e)))
            pg.on('console', lambda c: errs.append('console: ' + c.text) if c.type == 'error' else None)
            await pg.goto(uri)
            await pg.wait_for_timeout(900)
            for t in tabs:
                await pg.click('#tab-' + t)
                await pg.wait_for_timeout(320)
                ov = await pg.evaluate(
                    'document.documentElement.scrollWidth - document.documentElement.clientWidth')
                if ov > 0:
                    bad(f'{w}px {scheme}: {t} overflows by {ov}px')
                empty = await pg.evaluate(
                    f'''(() => {{const s=document.querySelector('#panel-{t}');
                        return [...s.querySelectorAll('svg.plot')]
                          .filter(x => x.childElementCount === 0).length;}})()''')
                if empty:
                    bad(f'{w}px {scheme}: {t} has {empty} empty chart(s)')
            if errs:
                bad(f'{w}px {scheme}: JS errors {errs[:3]}')
            else:
                good(f'{w}px {scheme}: all {len(tabs)} tabs render clean')
            await pg.close()

        # ---------------- filters actually change the numbers ----------------
        print('\n[B] filters re-scope every view')
        pg = await b.new_page(viewport={'width': 1320, 'height': 1000},
                              accept_downloads=True)
        errs = []
        pg.on('pageerror', lambda e: errs.append(str(e)))
        await pg.goto(uri)
        await pg.wait_for_timeout(800)
        base = await pg.evaluate("document.querySelector('#tiles .value').textContent")
        await pg.select_option('#f-region', label='Coastal')
        await pg.wait_for_timeout(500)
        coastal = await pg.evaluate("document.querySelector('#tiles .value').textContent")
        await pg.select_option('#f-region', index=0)
        await pg.select_option('#f-from', label='2015–16')
        await pg.wait_for_timeout(500)
        recent = await pg.evaluate("document.querySelector('#tiles .value').textContent")
        if base == coastal or base == recent:
            bad(f'filters did not change the headline stat ({base} / {coastal} / {recent})')
        else:
            good(f'region and season filters both re-scope ({base} → {coastal} / {recent})')
        await pg.click('#f-reset')
        await pg.wait_for_timeout(400)
        reset = await pg.evaluate("document.querySelector('#tiles .value').textContent")
        if reset != base:
            bad(f'reset did not restore the baseline ({reset} vs {base})')
        else:
            good('reset restores the baseline')

        # species chips
        await pg.click('#f-species .chip:nth-child(2)')
        await pg.wait_for_timeout(400)
        after = await pg.evaluate("document.querySelector('#tiles .value').textContent")
        if after == base:
            bad('toggling a species chip changed nothing')
        else:
            good(f'species chip re-scopes ({base} → {after})')
        await pg.click('#f-reset')
        await pg.wait_for_timeout(300)

        # ---------------- exports ----------------
        print('\n[C] exports download and open')
        await pg.click('#tab-data')
        await pg.wait_for_timeout(700)

        async def grab(label, scope_idx=None):
            if scope_idx is not None:
                await pg.click(f'#exp-scope .chip:nth-child({scope_idx})')
                await pg.wait_for_timeout(400)
            async with pg.expect_download(timeout=60000) as dl:
                await pg.click(f'.exp:has-text("{label}")')
            d = await dl.value
            path = DL / d.suggested_filename
            await d.save_as(path)
            return path

        # XLSX
        try:
            f = await grab('Excel workbook')
            from openpyxl import load_workbook
            wb = load_workbook(f)
            ws = wb[wb.sheetnames[0]]
            rows = ws.max_row - 1
            numeric = sum(1 for c in ws[2] if isinstance(c.value, (int, float)))
            if rows < 100 or numeric < 4:
                bad(f'xlsx thin: {rows} rows, {numeric} numeric cells in row 2')
            elif 'About' not in wb.sheetnames:
                bad('xlsx missing the About/provenance sheet')
            else:
                good(f'xlsx opens in openpyxl: {rows:,} rows, {ws.max_column} cols, '
                     f'{numeric} numeric cells, sheets {wb.sheetnames}')
        except Exception as e:
            bad(f'xlsx: {type(e).__name__}: {e}')

        # DOCX
        try:
            f = await grab('Word document')
            import docx
            doc = docx.Document(f)
            t = doc.tables[0] if doc.tables else None
            if not t or len(t.rows) < 5:
                bad('docx has no usable table')
            else:
                good(f'docx opens in python-docx: {len(doc.paragraphs)} paragraphs, '
                     f'table {len(t.rows)}×{len(t.columns)}')
        except Exception as e:
            bad(f'docx: {type(e).__name__}: {e}')

        # CSV
        try:
            f = await grab('Comma-separated')
            import csv as _csv
            txt = f.read_text('utf-8-sig')
            rows = list(_csv.reader(txt.splitlines()))
            if len(rows) < 100:
                bad(f'csv only {len(rows)} rows')
            else:
                good(f'csv parses: {len(rows):,} rows, {len(rows[0])} columns')
        except Exception as e:
            bad(f'csv: {type(e).__name__}: {e}')

        # GeoJSON
        try:
            f = await grab('GeoJSON')
            g = json.loads(f.read_text())
            feats = g['features']
            oob = [x for x in feats
                   if not (-125 < x['geometry']['coordinates'][0] < -116
                           and 45 < x['geometry']['coordinates'][1] < 49.5)]
            if not feats:
                bad('geojson has no features')
            elif oob:
                bad(f'{len(oob)} geojson points outside Washington')
            else:
                good(f'geojson valid: {len(feats)} points, all inside WA, '
                     f'{len(feats[0]["properties"])} properties each')
        except Exception as e:
            bad(f'geojson: {type(e).__name__}: {e}')

        # KML
        try:
            f = await grab('Google Earth')
            root = ET.fromstring(f.read_text())
            ns = '{http://www.opengis.net/kml/2.2}'
            marks = root.findall(f'.//{ns}Placemark')
            if not marks:
                bad('kml has no placemarks')
            else:
                good(f'kml is well-formed XML: {len(marks)} placemarks')
        except Exception as e:
            bad(f'kml: {type(e).__name__}: {e}')

        # every scope produces a working workbook
        for i, name in enumerate(['Every row', 'By season', 'By facility', 'Locations'], 1):
            try:
                f = await grab('Excel workbook', scope_idx=i)
                from openpyxl import load_workbook
                wb = load_workbook(f)
                ws = wb[wb.sheetnames[0]]
                if ws.max_row < 3:
                    bad(f'scope "{name}" produced an empty sheet')
                else:
                    good(f'scope "{name}": {ws.max_row - 1:,} rows × {ws.max_column} cols')
            except Exception as e:
                bad(f'scope "{name}": {type(e).__name__}: {e}')

        # ---------------- interactions ----------------
        print('\n[D] interactions')
        await pg.click('#tab-facility')
        await pg.wait_for_timeout(600)
        n0 = await pg.evaluate("document.querySelectorAll('#cmp-selected .chip').length")
        await pg.fill('#cmp-q', 'issaquah')
        await pg.wait_for_timeout(300)
        await pg.click('#cmp-list .opt')
        await pg.wait_for_timeout(400)
        n1 = await pg.evaluate("document.querySelectorAll('#cmp-selected .chip').length")
        if n1 != n0 + 1:
            bad(f'compare picker did not add a hatchery ({n0} → {n1})')
        else:
            good(f'compare picker adds ({n0} → {n1})')
        await pg.click('#cmp-selected .chip')
        await pg.wait_for_timeout(300)
        n2 = await pg.evaluate("document.querySelectorAll('#cmp-selected .chip').length")
        if n2 != n1 - 1:
            bad('removing a compare chip did nothing')
        else:
            good('compare chip removes')

        await pg.click('#tab-map')
        await pg.wait_for_timeout(700)
        pts = await pg.evaluate("document.querySelectorAll('#c-map .site').length")
        land = await pg.evaluate("document.querySelectorAll('#c-map .land').length")
        if pts < 50 or land < 1:
            bad(f'map thin: {pts} points, {land} land paths')
        else:
            good(f'map draws {pts} facilities on {land} land polygons')

        # tooltip appears on hover (scroll it into view first — mouse.move is
        # viewport-relative, so an off-screen mark can never be hovered)
        await pg.evaluate("document.querySelector('#c-map').scrollIntoView({block:'center'})")
        await pg.wait_for_timeout(400)
        box = await pg.evaluate("""(() => {const c=document.querySelector('#c-map .site-hit');
            const r=c.getBoundingClientRect(); return {x:r.x+r.width/2, y:r.y+r.height/2};})()""")
        await pg.mouse.move(box['x'], box['y'])
        await pg.wait_for_timeout(300)
        vis = await pg.evaluate("getComputedStyle(document.getElementById('tip')).opacity")
        if float(vis) < 0.9:
            bad('map tooltip did not appear on hover')
        else:
            good('map tooltip appears on hover')

        # table twins
        await pg.click('#tab-overview')
        await pg.wait_for_timeout(500)
        await pg.click('.tbl-toggle[data-table="t-trend"]')
        await pg.wait_for_timeout(300)
        tr = await pg.evaluate("document.querySelectorAll('#t-trend tbody tr').length")
        if tr < 5:
            bad(f'table twin empty ({tr} rows)')
        else:
            good(f'table twin renders {tr} rows')

        # ---------------- print ----------------
        print('\n[E] print layout')
        await pg.emulate_media(media='print')
        await pg.wait_for_timeout(400)
        hidden = await pg.evaluate("""(() => {
            const q = s => { const e=document.querySelector(s);
              return e ? getComputedStyle(e).display : 'missing'; };
            return {filters:q('.filters'), tabs:q('nav.tabs'), card:q('.card')};})()""")
        if hidden['filters'] != 'none' or hidden['tabs'] != 'none':
            bad(f'print still shows controls: {hidden}')
        else:
            good(f'print hides controls, keeps cards ({hidden["card"]})')
        pdf = DL / 'print_test.pdf'
        await pg.pdf(path=str(pdf), format='Letter', print_background=True)
        if pdf.stat().st_size < 20000:
            bad(f'printed PDF suspiciously small ({pdf.stat().st_size} bytes)')
        else:
            good(f'PDF renders: {pdf.stat().st_size // 1024} KB')
        await pg.emulate_media(media='screen')

        if errs:
            bad(f'JS errors during interaction: {errs[:3]}')
        await b.close()

    print('\n' + '=' * 62)
    print(f'UI TESTS: {"ALL PASSED" if not FAILS else str(len(FAILS)) + " FAILURE(S)"}')
    for f in FAILS:
        print('  ✗', f)
    print('=' * 62)
    return 1 if FAILS else 0


if __name__ == '__main__':
    sys.exit(asyncio.run(main()))
